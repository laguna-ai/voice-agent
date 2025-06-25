from langchain.document_loaders.csv_loader import CSVLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from io import BytesIO
from docx import Document as DocxDocument
from vectorstores.postgres import create_vectorstore
from MSAL.search import get_all_files_info, get_file_content
from langchain_core.documents import Document
from pypdf import PdfReader
from docx.opc.exceptions import PackageNotFoundError  # Para DOCX corruptos
from pypdf.errors import PdfReadError  # Para PDFs inválidos


def create_sharepoint_index():
    """Crea una base de conocimientos desde SharePoint con DOCX y PDF en memoria"""

    # 1. Obtener y filtrar archivos
    valid_files = [
        item
        for item in get_all_files_info()
        if item["type"] == "file" and item["name"].lower().endswith((".pdf", ".docx"))
    ]

    all_docs = []

    # 2. Procesamiento común en función
    for file_info in valid_files:
        try:
            file_bytes = get_file_content(file_info["id"])
            base_metadata = {
                "source": file_info["path"],
                "file_name": file_info["name"],
            }

            # Procesamiento específico por tipo
            if file_info["name"].lower().endswith(".docx"):
                doc = DocxDocument(BytesIO(file_bytes))
                # Texto normal de párrafos
                text_content = "\n".join(p.text for p in doc.paragraphs)
                all_docs.append(
                    Document(page_content=text_content, metadata=base_metadata)
                )

                # Procesamiento de tablas
                for table in doc.tables:
                    headers = [cell.text.strip() for cell in table.rows[0].cells]

                    for row in table.rows[1:]:  # Saltar la fila de encabezados
                        row_data = []
                        for i, cell in enumerate(row.cells):
                            if i < len(headers):  # Prevenir index error
                                row_data.append(f"{headers[i]}: {cell.text.strip()}")

                        if row_data:  # Solo añadir filas con contenido
                            table_content = "\n".join(row_data)
                            all_docs.append(
                                Document(
                                    page_content=table_content,
                                    metadata={**base_metadata, "tipo": "tabla"},
                                )
                            )
            elif file_info["name"].lower().endswith(".pdf"):
                pdf = PdfReader(BytesIO(file_bytes))
                for page_num, page in enumerate(pdf.pages, 1):
                    all_docs.append(
                        Document(
                            page_content=page.extract_text(),
                            metadata={
                                **base_metadata,
                                "page": page_num,
                                "total_pages": len(pdf.pages),
                            },
                        )
                    )
        except (PackageNotFoundError, ValueError) as e:
            print(f"Archivo Word inválido/corrupto: {file_info['name']} - {str(e)}")
            continue
        except PdfReadError as e:
            print(f"Error leyendo PDF: {file_info['name']} - {str(e)}")
            continue

    # 3. Dividir y formatear documentos
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1200, chunk_overlap=250)
    structured_docs = [d for d in all_docs if d.metadata.get("tipo") == "tabla"]
    unstructured_docs = [d for d in all_docs if d.metadata.get("tipo") != "tabla"]
    unstructured_splits = text_splitter.split_documents(unstructured_docs)
    all_splits = structured_docs + unstructured_splits

    # Añadir metadatos al contenido
    for split in all_splits:
        metadata = split.metadata
        # tomar solo nombre de carpeta para la ruta (lo que está antes de /)
        metadata["categoria"] = metadata["source"].split("/")[0]
        split.page_content = f"{metadata['source']}:\n {split.page_content}"

    print(f"Total de splits generados: {len(all_splits)}")
    # Mostrar los primeros 200 caracteres del primer split
    print("Primer split:", all_splits[0])
    # 6. Crear vectorstore en chunks de 500 splits
    # (para evitar problemas )
    chunk_size = 500
    for i in range(0, len(all_splits), chunk_size):
        chunk = all_splits[i : i + chunk_size]
        create_vectorstore(chunk)


def create_csv_index(config_archivos_csv):
    """
    Crea una base de conocimientos en postgres a partir de una lista de archivos CSV con configuraciones específicas.

    :param config_archivos_csv: Lista de diccionarios con configuraciones de archivos CSV.
                                Cada diccionario debe tener las claves: 'file_path', 'encoding', y 'delimiter'.
    """
    # Cargar y combinar todos los datos de los archivos CSV
    datos_combinados = []
    for config in config_archivos_csv:
        loader = CSVLoader(
            file_path=config["file_path"],
            encoding=config["encoding"],
            csv_args={"delimiter": config["delimiter"]},
            metadata_columns=["categoria"],
        )
        datos = loader.load()

        for dato in datos:
            dato.page_content += (
                f" Referencia: {config['file_path'].split('/')[1].split('.')[0]}"
            )

        datos_combinados.extend(datos)
    # ver datos
    print("TOTAL DATOS CSV:", len(datos_combinados))
    # Crear embeddings
    create_vectorstore(datos_combinados)


# Ejemplo de uso
Dir = "csv_data"
conf_archivos_csv = [
    {"file_path": f"{Dir}/extra.csv", "encoding": "UTF-8", "delimiter": ","},
]

# {"file_path": f"{Dir}/faqs.csv", "encoding": "UTF-8", "delimiter": ","},
# {"file_path": f"{Dir}/cuestionario formativo unidad 1 Diseño de Autor.csv", "encoding": "UTF-8", "delimiter": ","}
# , {"file_path": f"{Dir}/cuestionario formativo unidad 2 Diseño de Autor.csv", "encoding": "UTF-8", "delimiter": ","}
# , {"file_path": f"{Dir}/cuestionario formativo unidad 3 Diseño de Autor.csv", "encoding": "UTF-8", "delimiter": ","}
# , {"file_path": f"{Dir}/cuestionario formativo unidad 4 Diseño de Autor.csv", "encoding": "UTF-8", "delimiter": ","}
# , {"file_path": f"{Dir}/cuestionario sumativo unidad 5 Diseño de Autor.csv", "encoding": "UTF-8", "delimiter": ","}
# , {"file_path": f"{Dir}/contenido unidad 1 Diseño de Autor.csv", "encoding": "UTF-8", "delimiter": ","}
# , {"file_path": f"{Dir}/contenido unidad 2 Diseño de Autor.csv", "encoding": "UTF-8", "delimiter": ","}
# , {"file_path": f"{Dir}/contenido unidad 3 Diseño de Autor.csv", "encoding": "UTF-8", "delimiter": ","}
# , {"file_path": f"{Dir}/contenido unidad 4 Diseño de Autor.csv", "encoding": "UTF-8", "delimiter": ","}
# , {"file_path": f"{Dir}/contenido unidad 5 Diseño de Autor.csv", "encoding": "UTF-8", "delimiter": ","}
# , {"file_path": f"{Dir}/contenido unidad 1 Diseño Instruccional.csv", "encoding": "UTF-8", "delimiter": ","}
# , {"file_path": f"{Dir}/contenido unidad 2 Diseño Instruccional.csv", "encoding": "UTF-8", "delimiter": ","}
# , {"file_path": f"{Dir}/contenido unidad 3 Diseño Instruccional.csv", "encoding": "UTF-8", "delimiter": ","}
# , {"file_path": f"{Dir}/contenido unidad 4 Diseño Instruccional.csv", "encoding": "UTF-8", "delimiter": ","}
# , {"file_path": f"{Dir}/contenido unidad 5 Diseño Instruccional.csv", "encoding": "UTF-8", "delimiter": ","}
# , {"file_path": f"{Dir}/cuestionario unidad 1 Diseño Instruccional.csv", "encoding": "UTF-8", "delimiter": ","}
# , {"file_path": f"{Dir}/cuestionario unidad 4 Diseño Instruccional.csv", "encoding": "UTF-8", "delimiter": ","}
# , {"file_path": f"{Dir}/programa Diseño de Autor.csv", "encoding": "UTF-8", "delimiter": ","}
# , {"file_path": f"{Dir}/programa Diseño Instruccional.csv", "encoding": "UTF-8", "delimiter": ","}

# create_sharepoint_index()
create_csv_index(conf_archivos_csv)
